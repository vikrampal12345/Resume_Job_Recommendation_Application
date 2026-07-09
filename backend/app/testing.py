from docx import Document

doc = Document("app\\test_files\\Milano-Basic-_-Simple-Resume-Template-Green.docx")

# print("Paragraphs:", len(doc.paragraphs))

# for para in doc.paragraphs:
#     print(para.text)

# from docx import Document

# doc = Document("app/test_files/sample_resume.docx")

print("Tables:", len(doc.tables))

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            print(cell.text)