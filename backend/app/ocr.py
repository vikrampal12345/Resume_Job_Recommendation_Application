import os
import fitz
# import easyocr
from docx import Document
from pymupdf import FileDataError

# ===========================================
# Load EasyOCR Once
# ===========================================

# reader = easyocr.Reader(['en'], gpu=False)
# reader = None

# ===========================================
# PDF Extraction
# ===========================================

def extract_pdf_text(file_path):
   
    
 


    try:
        pdf = fitz.open(file_path)

        text = ""

        for page in pdf:
            text += page.get_text()

        pdf.close()

        if not text.strip():
            raise ValueError("No text found in the PDF.")

        return text

    except FileDataError:
        raise ValueError("Invalid or corrupted PDF file.")

    except Exception as e:
        raise ValueError(f"Unable to read PDF: {e}")


# ===========================================
# DOCX Extraction
# ===========================================

def extract_docx_text(file_path):

    try:
        doc = Document(file_path)

        text = []

        # Paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text.strip())

        # Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text.strip())

        if not text:
            raise ValueError("No text found in the DOCX file.")

        return "\n".join(text)

    except Exception as e:
        raise ValueError(f"Unable to read DOCX: {e}")


# ===========================================
# Image OCR
# ===========================================

# def extract_image_text(file_path):
    
#     global reader

#     if reader is None:

#         print("Loading EasyOCR...")

#         reader = easyocr.Reader(['en'], gpu=False)

#     try:

#         result = reader.readtext(file_path)

#         text = ""

#         for item in result:
#             text += item[1] + " "

#         if not text.strip():
#             raise ValueError("No text detected in image.")

#         return text

#     except Exception as e:
#         raise ValueError(f"Unable to read image: {e}")


# ===========================================
# Main Function
# ===========================================

def extract_resume_text(file_path):

    # File exists?
    if not os.path.exists(file_path):
        raise FileNotFoundError("File does not exist.")

    # macOS hidden file
    if os.path.basename(file_path).startswith("._"):
        raise ValueError(
            "Invalid macOS metadata file. Please select the actual resume file."
        )

    extension = os.path.splitext(file_path)[1].lower()

    if extension == ".pdf":
        return extract_pdf_text(file_path)

    elif extension == ".docx":
        return extract_docx_text(file_path)

    elif extension in [".png", ".jpg", ".jpeg"]:
        return extract_image_text(file_path)

    else:
        raise ValueError(
            "Unsupported file format. Supported formats: PDF, DOCX."
        )


# ===========================================
# Testing
# ===========================================

if __name__ == "__main__":

    file_path = "app/test_files/Milano-Resume-Template-Black.pdf"

    try:

        text = extract_resume_text(file_path)

        print("=" * 50)
        print("Resume Text")
        print("=" * 50)

        print(text[:1500])

    except Exception as e:

        print("=" * 50)
        print("ERROR")
        print("=" * 50)

        print(e)